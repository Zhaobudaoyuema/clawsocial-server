"""Engine B: python-docx fallback with same ReplacementPlan."""
from pathlib import Path

from docx import Document

from app.deid.engine.plan import ReplacementPlan


def run_docx_fallback(input_path: Path, output_path: Path, plan: ReplacementPlan) -> tuple[int, dict]:
    doc = Document(str(input_path))
    total = 0
    for para in doc.paragraphs:
        text = para.text
        if not text.strip():
            continue
        new_text, cnt = plan.apply_to_text(text)
        if cnt:
            para.clear()
            para.add_run(new_text)
            total += cnt
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    if not text.strip():
                        continue
                    new_text, cnt = plan.apply_to_text(text)
                    if cnt:
                        para.clear()
                        para.add_run(new_text)
                        total += cnt
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    coverage = {"body": total, "header": 0, "footer": 0, "footnote": 0, "textbox": 0}
    return total, coverage
