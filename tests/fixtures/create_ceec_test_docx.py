"""Generate CEEC-themed de-id test docx (simulated audit workpaper)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

OUT = Path(__file__).parent / "ceec_audit_test.docx"


def _set_run_font(run, name="宋体", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run_font(run, "黑体", 14 if level == 1 else 12, bold=True)
    return p


def _para(doc, text, bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, bold=bold)
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                _set_run_font(run, "黑体", 10, bold=True)
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    _set_run_font(run, size=10)
    return t


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    # 封面
    _para(doc, "中国能源建设股份有限公司", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "2024年度财务报表审计底稿（测试稿）", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "—— 合并范围关联方及往来函证工作底稿 ——", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    _para(doc, "编制单位：某会计师事务所（模拟）")
    _para(doc, "被审计单位：中国能源建设股份有限公司（股票代码：601868.SH / 3996.HK）")
    _para(doc, "报告期间：2024年1月1日至2024年12月31日")
    _para(doc, "密级：内部资料 · 仅供脱敏工具测试")
    doc.add_page_break()

    _heading(doc, "一、审计目标与范围")
    _para(
        doc,
        "本底稿用于记录对中国能源建设股份有限公司（以下简称“能建股份”或“中国能建”）"
        "及其合并范围内主要子公司 2024 年度关联方交易、往来余额的审计程序与结论。"
        "母公司中国能源建设集团有限公司为最终控股方。",
    )
    _para(
        doc,
        "合并范围包括：中国葛洲坝集团股份有限公司、电力规划总院有限公司、"
        "中国电力工程顾问集团有限公司、中能建国际建设集团有限公司、"
        "中国能源建设集团投资有限公司、中国能源建设集团财务有限公司等。",
    )

    _heading(doc, "二、主要财务数据摘录（金额应保留）", level=2)
    _para(doc, "以下数据摘自模拟合并利润表及资产负债表，单位：人民币千元。")
    _table(
        doc,
        ["项目", "2024年度", "2023年度", "变动率"],
        [
            ["营业收入", "436,713,000", "406,032,000", "7.56%"],
            ["营业成本", "389,245,000", "362,108,000", "7.49%"],
            ["毛利率", "10.87%", "10.82%", "+0.05pct"],
            ["净利润", "8,956,000", "8,412,000", "6.47%"],
            ["应收账款账面余额", "92,109,000", "85,234,000", "8.07%"],
            ["合同资产账面余额", "106,688,000", "98,456,000", "8.36%"],
        ],
    )
    doc.add_paragraph()

    _heading(doc, "三、关联方清单及交易发生额", level=2)
    _table(
        doc,
        ["关联方名称", "关联关系", "交易类型", "本期发生额（千元）", "期末余额（千元）"],
        [
            [
                "中国能源建设集团有限公司",
                "母公司",
                "资金拆借",
                "12,500,000",
                "3,200,000",
            ],
            [
                "中国葛洲坝集团股份有限公司",
                "子公司",
                "工程分包",
                "45,678,000",
                "8,901,000",
            ],
            [
                "电力规划总院有限公司",
                "同系子公司",
                "勘测设计服务",
                "2,345,000",
                "456,000",
            ],
            [
                "中能建国际建设集团有限公司",
                "同系子公司",
                "境外工程承包",
                "6,789,000",
                "1,234,000",
            ],
            [
                "中国能源建设集团财务有限公司",
                "同系子公司",
                "存款及贷款",
                "18,900,000",
                "5,670,000",
            ],
            [
                "中国能源建设集团华东建设投资有限公司",
                "同系子公司",
                "项目投资",
                "890,000",
                "234,000",
            ],
            [
                "中国能源建设集团广东火电工程有限公司",
                "同系子公司",
                "设备采购",
                "1,567,000",
                "345,000",
            ],
        ],
    )
    doc.add_paragraph()

    _heading(doc, "四、函证及审计程序记录", level=2)
    _para(
        doc,
        "1. 向葛洲坝股份发送询证函，确认期末应收工程款余额 8,901,000 千元，"
        "已收到回函相符。",
    )
    _para(
        doc,
        "2. 检查能建投资与能建租赁之间的融资租赁合同，利率 3.85%，"
        "符合市场公允价格区间。",
    )
    _para(
        doc,
        "3. 复核北京电建、浙江火电、江苏电建三公司对能建股份的分包结算单，"
        "抽样比例 65%，未发现重大异常。",
    )
    _para(
        doc,
        "4. 对中能建装备销售给母公司的专用设备执行截止测试，"
        "收入确认时点与验收单一致。",
    )

    _heading(doc, "五、区域子公司补充说明（正则测试用）", level=2)
    _para(doc, "以下主体名称用于测试区域公司识别规则：")
    bullets = [
        "华东区域公司与中国能建签署战略合作协议，新签合同额 23.6 亿元。",
        "华北区域公司参与雄安新区综合能源项目，投资比例 35%。",
        "西北区域公司承建甘肃酒泉风电场 EPC 总承包工程。",
        "华南区域公司与广东火电联合投标海上风电项目。",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for run in p.runs:
            _set_run_font(run)

    _heading(doc, "六、需人工补充的漏网主体（测试手动添加）", level=2)
    _para(
        doc,
        "审计师现场访谈获悉，被审计单位与「中能建氢能源有限公司」存在技术服务往来，"
        "该主体尚未纳入预设词库，需在脱敏流程中手动补充。本期发生额 128,000 千元。",
    )
    _para(
        doc,
        "另发现与「某省能源建设投资有限公司」（非上市体系外参股公司）存在少量代垫款，"
        "余额 45,000 千元，建议项目组评估是否纳入脱敏范围。",
    )

    _heading(doc, "七、审计结论", level=2)
    _para(
        doc,
        "经实施上述程序，我们认为中国能源建设股份有限公司 2024 年度关联方交易披露"
        "在所有重大方面符合企业会计准则第 36 号——关联方披露的要求。"
        "营业收入 4,367.13 亿元中，工程建设业务收入占比 79.46%，"
        "毛利率 10.87%，与行业平均水平无重大偏离。",
    )
    _para(doc, "编制：张某某    复核：李某某    日期：2025-03-15")

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
