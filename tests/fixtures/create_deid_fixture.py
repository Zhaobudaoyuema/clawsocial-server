"""Generate minimal test docx with SPIC terms."""
from pathlib import Path

from docx import Document

OUT = Path(__file__).parent / "spic_sample.docx"


def main():
    doc = Document()
    doc.add_paragraph("审计底稿：国家电投集团下属单位往来。")
    doc.add_paragraph("涉及国家 电投与电投产融的关联交易。")
    doc.add_paragraph("应收账款余额 1,234,567.89 元，毛利率 12.5%。")
    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
