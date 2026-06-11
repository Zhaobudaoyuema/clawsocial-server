"""Tests for de-identification API and engine."""
import json
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.deid.engine.plan import normalize_for_match, build_plan_from_job_entities
from scripts.seed_deid_spic import seed_db

FIXTURE = Path(__file__).parent / "fixtures" / "spic_sample.docx"
CEEC_FIXTURE = Path(__file__).parent / "fixtures" / "ceec_audit_test.docx"


def _ensure_fixture():
    if not FIXTURE.exists():
        doc = Document()
        doc.add_paragraph("审计底稿：国家电投集团下属单位往来。")
        doc.add_paragraph("涉及国家 电投与电投产融的关联交易。")
        doc.add_paragraph("应收账款余额 1,234,567.89 元。")
        FIXTURE.parent.mkdir(parents=True, exist_ok=True)
        doc.save(FIXTURE)


def _entity_ids(client: TestClient, job_id: int) -> list[int]:
    ents = client.get(f"/api/deid/jobs/{job_id}/entities").json()
    return [e["id"] for e in ents if not e.get("is_excluded")]


def _complete_program_scan(client: TestClient, job_id: int) -> None:
    """Run program scan and ack before confirm/run."""
    r = client.post(f"/api/deid/jobs/{job_id}/program-scan/run")
    assert r.status_code == 200, r.text
    r = client.post(f"/api/deid/jobs/{job_id}/program-scan/confirm")
    assert r.status_code == 200, r.text


def test_normalize_for_match():
    assert normalize_for_match("国家 电投") == normalize_for_match("国家电投")


def test_extract_sample_text_from_markdown(tmp_path):
    from app.deid.engine.markdown_pipeline import extract_sample_text

    md_path = tmp_path / "source.md"
    md_path.write_text("中国能源建设股份有限公司审计报告\n", encoding="utf-8")
    sample = extract_sample_text(md_path)
    assert "中国能源建设股份有限公司" in sample


def test_plan_matching():
    entities = [
        {
            "canonical_name": "国家电投",
            "entity_type": "company",
            "placeholder": "[公司_1]",
            "source": "remembered",
            "is_excluded": False,
            "aliases": ["国家电投", "国家 电投"],
        }
    ]
    plan = build_plan_from_job_entities(entities, [], [])
    out, cnt = plan.apply_to_text("国家 电投与电投产融")
    assert cnt >= 1
    assert "[公司_1]" in out


def test_seed_idempotent(client: TestClient, db, seeded_db):
    packs = client.get("/api/deid/packs")
    assert packs.status_code == 200
    codes = {p["code"] for p in packs.json()}
    assert "spic" in codes
    assert "ceec" in codes
    assert "general_finance" in codes


def test_job_scan_run_flow(client: TestClient, db, seeded_db):
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    assert r.status_code == 200
    job = r.json()
    job_id = job["id"]
    assert job.get("file_type") == "markdown"

    r = client.get(f"/api/deid/jobs/{job_id}/source-markdown")
    assert r.status_code == 200
    md = r.json()
    assert md["source_format"] == "docx"
    assert md["source_format_label"] == "Word"
    assert md["file_type"] == "markdown"
    assert "content" in md
    assert md["stats"]["char_count"] >= 50

    r = client.post(f"/api/deid/jobs/{job_id}/scan")
    assert r.status_code == 200
    data = r.json()
    entities = data.get("entities") or data
    if isinstance(entities, dict):
        entities = entities.get("entities", [])
    names = " ".join(e["canonical_name"] for e in entities)
    assert "电投" in names or "国家" in names

    ids = _entity_ids(client, job_id)
    assert ids
    _complete_program_scan(client, job_id)
    r = client.post(f"/api/deid/jobs/{job_id}/run", json={"entity_ids": ids})
    assert r.status_code == 200
    result = r.json()
    assert result["status"] == "done"

    r = client.get(f"/api/deid/jobs/{job_id}/export?override_ack=true")
    assert r.status_code == 200
    ct = r.headers.get("content-type", "")
    assert "markdown" in ct or "text/plain" in ct or "octet-stream" in ct
    cd = r.headers.get("content-disposition", "")
    assert "deid_" in cd
    assert "_desensitized.md" in cd


def _ensure_ceec_fixture():
    if not CEEC_FIXTURE.exists():
        from tests.fixtures.create_ceec_test_docx import main as gen

        gen()


def test_ceec_fixture_scan(client: TestClient, db, seeded_db):
    """Regression: canonical name duplicated in alias rows must not 500 on scan."""
    _ensure_ceec_fixture()
    with open(CEEC_FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "ceec_audit_test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert r.status_code == 200
    job_id = r.json()["id"]
    r = client.post(f"/api/deid/jobs/{job_id}/scan")
    assert r.status_code == 200, r.text
    data = r.json()
    entities = data.get("entities") or data
    if isinstance(entities, dict):
        entities = entities.get("entities", [])
    names = " ".join(e["canonical_name"] for e in entities)
    assert "能建" in names or "葛洲坝" in names
    summary = data.get("scan_summary") or {}
    assert summary.get("llm_skipped") in ("worker_offline", "llm_disabled", None) or summary.get("llm_hits", 0) >= 0


def test_worker_status_offline(client: TestClient, monkeypatch):
    """Isolate from remote DEID relay configured on dev machines."""
    relay = client.app.state.worker_relay
    relay._base_url = ""
    relay._token = ""
    relay._status = None
    monkeypatch.setattr(
        "app.api.deid.get_worker_client",
        lambda app: app.state.worker_router,
    )
    client.app.state.worker_router._session = None
    r = client.get("/api/deid/worker/status")
    assert r.status_code == 200
    data = r.json()
    assert data["online"] is False
    assert data["state"] == "offline"


def test_start_scan_without_worker(client: TestClient, db, seeded_db):
    """Async start + poll path; use_worker=false skips LLM queue."""
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"use_worker": "false"},
        )
    assert r.status_code == 200
    job_id = r.json()["id"]
    assert r.json()["use_worker"] is False

    r = client.post(f"/api/deid/jobs/{job_id}/start")
    assert r.status_code == 200
    assert r.json()["status"] in ("scanning", "queued")

    import time

    for _ in range(30):
        r = client.get(f"/api/deid/jobs/{job_id}")
        assert r.status_code == 200
        body = r.json()
        if body["status"] == "scanned":
            assert body.get("progress", {}).get("percent") == 100
            break
        time.sleep(0.2)
    else:
        raise AssertionError("scan did not complete")


def test_scan_with_mock_llm(client: TestClient, db, seeded_db, monkeypatch):
    _ensure_ceec_fixture()
    import app.deid.service as service_mod
    from app.deid.discovery.entity_rescan import InitialScanResult
    from app.deid.discovery.merge import MergedEntity

    async def fake_initial_scan(*args, **kwargs):
        entities = [
            MergedEntity(
                canonical_name="中能建氢能源有限公司",
                entity_type="company",
                source="llm",
                aliases=["中能建氢能源有限公司"],
                hit_count=2,
                confidence=0.92,
            )
        ]
        return InitialScanResult(entities=entities, llm_result=None)

    monkeypatch.setattr(service_mod, "run_initial_scan", fake_initial_scan)
    monkeypatch.setattr(service_mod, "job_needs_worker_queue", lambda db, job, router: True)
    monkeypatch.setenv("DEID_LLM_ENABLED", "1")

    with open(CEEC_FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "ceec_audit_test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    job_id = r.json()["id"]
    r = client.post(f"/api/deid/jobs/{job_id}/scan")
    assert r.status_code == 200
    data = r.json()
    entities = data["entities"]
    names = " ".join(e["canonical_name"] for e in entities)
    assert "氢能源" in names
    assert data["scan_summary"]["llm_hits"] >= 1


def test_llm_entities_auto_saved_to_library(client: TestClient, db, seeded_db, monkeypatch):
    """LLM entities are saved to library by default when deid runs."""
    _ensure_ceec_fixture()
    import app.deid.service as service_mod
    from app.deid.discovery.entity_rescan import InitialScanResult
    from app.deid.discovery.merge import MergedEntity

    async def fake_initial_scan(*args, **kwargs):
        entities = [
            MergedEntity(
                canonical_name="LLM测试主体有限公司",
                entity_type="company",
                source="llm",
                aliases=["LLM测试主体有限公司", "测试主体"],
                hit_count=1,
                confidence=0.9,
            )
        ]
        return InitialScanResult(entities=entities, llm_result=None)

    monkeypatch.setattr(service_mod, "run_initial_scan", fake_initial_scan)
    monkeypatch.setattr(service_mod, "job_needs_worker_queue", lambda db, job, router: True)
    monkeypatch.setenv("DEID_LLM_ENABLED", "1")

    with open(CEEC_FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "ceec_audit_test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    ents = client.get(f"/api/deid/jobs/{job_id}/entities").json()
    ids = [e["id"] for e in ents if not e.get("is_excluded")]
    _complete_program_scan(client, job_id)
    client.post(
        f"/api/deid/jobs/{job_id}/run",
        json={"entity_ids": ids},
    )

    lib = client.get("/api/deid/entities").json()
    match = [e for e in lib if "LLM测试主体" in e["canonical_name"]]
    assert len(match) == 1
    assert match[0]["source"] == "llm"


def test_confirm_with_remember_ids_no_duplicate_alias(client: TestClient, db, seeded_db, monkeypatch):
    """confirm 时 auto-save + remember_ids 重叠不应触发别名唯一约束冲突。"""
    _ensure_ceec_fixture()
    import app.deid.service as service_mod
    from app.deid.discovery.entity_rescan import InitialScanResult
    from app.deid.discovery.merge import MergedEntity

    async def fake_initial_scan(*args, **kwargs):
        return InitialScanResult(
            entities=[
                MergedEntity(
                    canonical_name="重复别名测试公司",
                    entity_type="company",
                    source="llm",
                    aliases=["重复别名测试公司", "测试简称"],
                    hit_count=1,
                )
            ],
            llm_result=None,
        )

    monkeypatch.setattr(service_mod, "run_initial_scan", fake_initial_scan)
    monkeypatch.setattr(service_mod, "job_needs_worker_queue", lambda db, job, router: True)
    monkeypatch.setenv("DEID_LLM_ENABLED", "1")

    with open(CEEC_FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "ceec_audit_test.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    ents = client.get(f"/api/deid/jobs/{job_id}/entities").json()
    ids = [e["id"] for e in ents if not e.get("is_excluded")]
    _complete_program_scan(client, job_id)
    r = client.post(
        f"/api/deid/jobs/{job_id}/confirm",
        json={"entity_ids": ids, "remember_ids": ids},
    )
    assert r.status_code == 200, r.text
    assert r.json()["status"] == "done"


def test_list_jobs_includes_incomplete(client: TestClient, db, seeded_db):
    """Sidebar lists incomplete jobs plus 8h done jobs."""
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "spic_sample.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    draft_id = r.json()["id"]
    listed = client.get("/api/deid/jobs").json()
    assert any(j["id"] == draft_id for j in listed)

    client.post(f"/api/deid/jobs/{draft_id}/scan")
    ids = _entity_ids(client, draft_id)
    _complete_program_scan(client, draft_id)
    client.post(f"/api/deid/jobs/{draft_id}/run", json={"entity_ids": ids})
    listed = client.get("/api/deid/jobs").json()
    assert any(j["id"] == draft_id and j["status"] == "done" for j in listed)


def test_list_jobs_includes_archived(client: TestClient, db, seeded_db):
    from app.deid.service import archive_job_files
    from app.models_deid import DeidJob

    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "spic_sample.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"use_worker": "false"},
        )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    ids = _entity_ids(client, job_id)
    _complete_program_scan(client, job_id)
    client.post(f"/api/deid/jobs/{job_id}/run", json={"entity_ids": ids})
    job = db.get(DeidJob, job_id)
    archive_job_files(db, job)
    db.commit()
    listed = client.get("/api/deid/jobs").json()
    row = next(j for j in listed if j["id"] == job_id)
    assert row["status"] == "archived"
    assert row["rehydrate_available"] is True


def test_list_jobs_after_run(client: TestClient, db, seeded_db):
    """Regression: list_jobs must not crash on naive expires_at from MySQL."""
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={
                "file": (
                    "spic_sample.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    ids = _entity_ids(client, job_id)
    _complete_program_scan(client, job_id)
    client.post(f"/api/deid/jobs/{job_id}/run", json={"entity_ids": ids})
    r = client.get("/api/deid/jobs")
    assert r.status_code == 200
    assert any(j["id"] == job_id for j in r.json())


def test_manual_entity_and_library(client: TestClient, db, seeded_db):
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    r = client.post(
        f"/api/deid/jobs/{job_id}/entities",
        json={
            "canonical_name": "测试主体XYZ",
            "entity_type": "company",
            "aliases": ["测试主体XYZ"],
            "save_to_library": True,
        },
    )
    assert r.status_code == 200

    ents = client.get("/api/deid/entities").json()
    assert any("测试主体" in e["canonical_name"] for e in ents)


def test_remembered_source_label(client: TestClient, db, seeded_db):
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"use_worker": "false"},
        )
    job_id = r.json()["id"]
    r = client.post(f"/api/deid/jobs/{job_id}/scan")
    entities = r.json()["entities"]
    remembered = [e for e in entities if e["source"] == "remembered"]
    assert remembered
    assert remembered[0]["source_label"] == "已记住"


def test_entity_types_crud(client: TestClient, db, seeded_db):
    r = client.get("/api/deid/entity-types")
    assert r.status_code == 200
    codes = {t["code"] for t in r.json()}
    assert "company" in codes

    r = client.post(
        "/api/deid/entity-types",
        json={"code": "project", "label": "项目", "placeholder_prefix": "项目"},
    )
    assert r.status_code == 200
    assert r.json()["code"] == "project"

    r = client.patch(
        "/api/deid/entity-types/project",
        json={"label": "工程项目"},
    )
    assert r.status_code == 200
    assert r.json()["label"] == "工程项目"

    r = client.delete("/api/deid/entity-types/project")
    assert r.status_code == 200

    r = client.delete("/api/deid/entity-types/company")
    assert "不可删除" in r.text
    types = client.get("/api/deid/entity-types").json()
    assert any(t["code"] == "company" for t in types)


def test_delete_job(client: TestClient, db, seeded_db):
    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    job_id = r.json()["id"]
    r = client.delete(f"/api/deid/jobs/{job_id}")
    assert r.status_code == 200
    assert client.get(f"/api/deid/jobs/{job_id}").status_code == 404


def test_delete_job_with_global_experience(client: TestClient, db, seeded_db):
    from app.deid.discovery.experience_store import append_global_experience

    _ensure_fixture()
    with open(FIXTURE, "rb") as f:
        r = client.post(
            "/api/deid/jobs",
            files={"file": ("spic_sample.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
    job_id = r.json()["id"]
    append_global_experience(db, "表头简称需补全", source_job_id=job_id)
    db.commit()

    r = client.delete(f"/api/deid/jobs/{job_id}")
    assert r.status_code == 200
    assert client.get(f"/api/deid/jobs/{job_id}").status_code == 404
    from app.models_deid import DeidGlobalExperience

    row = db.query(DeidGlobalExperience).filter(DeidGlobalExperience.text == "表头简称需补全").one()
    assert row.source_job_id is None
