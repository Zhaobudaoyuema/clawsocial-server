"""Tests for program scan step (pre-confirm dry-run + auto-fix)."""
from pathlib import Path

from app.deid.discovery.program_scan import plan_fixes, simulate_run
from app.deid.engine.markdown_pipeline import residual_scan_text
from app.deid.engine.plan import build_plan_from_job_entities, normalize_for_match


def test_build_plan_includes_canonical_name():
    entities = [
        {
            "canonical_name": "中国能建",
            "entity_type": "company",
            "placeholder": "[公司_1]",
            "source": "llm",
            "is_excluded": False,
            "aliases": ["能建集团"],
        }
    ]
    plan = build_plan_from_job_entities(entities, [], [])
    plan.finalize()
    norms = {p.norm for p in plan._prepared}
    assert normalize_for_match("中国能建") in norms
    assert normalize_for_match("能建集团") in norms


def test_residual_scan_skips_placeholder_substrings():
    text = "近年来[公司_64]经营业务快速发展"
    entities = [
        {
            "canonical_name": "公司",
            "entity_type": "company",
            "placeholder": "[公司_1]",
            "source": "manual",
            "is_excluded": False,
            "aliases": ["公司"],
        }
    ]
    result = residual_scan_text(text, entities)
    assert result["passed"] is True


def test_c3_expand_alias_from_longer_form():
    source = "中国能源建设集团浙江火电建设有限公司承担本项目。"
    entities = [
        {
            "id": 1,
            "canonical_name": "中国能建",
            "entity_type": "company",
            "placeholder": "[公司_1]",
            "source": "llm",
            "is_excluded": False,
            "aliases": ["中国能建"],
        }
    ]
    before, fixes = simulate_run(source, entities, [], [], type_prefix_map={"company": "公司"})
    assert fixes
    assert any(f["action"] == "add_alias" for f in fixes)
    assert any("浙江火电" in f["text"] for f in fixes)


def test_c3_new_entity_when_unrelated():
    source = "中国建设银行股份有限公司-易方达基金出现在表中。"
    entities = [
        {
            "id": 1,
            "canonical_name": "中国能建",
            "entity_type": "company",
            "placeholder": "[公司_1]",
            "source": "llm",
            "is_excluded": False,
            "aliases": ["中国能建"],
        }
    ]
    preview = __import__(
        "app.deid.engine.preview", fromlist=["build_preview_text"]
    ).build_preview_text(source, entities, [], [], type_prefix_map={"company": "公司"})
    before = residual_scan_text(preview, entities)
    fixes = plan_fixes(source, preview, entities, before)
    assert any(f["action"] == "new_entity" for f in fixes)


def test_program_scan_run_and_confirm(client, db, seeded_db):
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "中国能建承担任务。中国能源建设集团浙江火电建设有限公司参与建设。" * 3
    )
    buf = __import__("io").BytesIO()
    doc.save(buf)
    buf.seek(0)

    r = client.post(
        "/api/deid/jobs",
        files={"file": ("program_scan_sample.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    client.post(f"/api/deid/jobs/{job_id}/semantic/skip")

    r = client.post(f"/api/deid/jobs/{job_id}/program-scan/run")
    assert r.status_code == 200
    data = r.json()
    assert data["status"] == "program_review"

    r = client.post(f"/api/deid/jobs/{job_id}/program-scan/confirm")
    assert r.status_code == 200
    assert r.json().get("program_scan_ack_at")

    ids = [e["id"] for e in client.get(f"/api/deid/jobs/{job_id}/entities").json() if not e.get("is_excluded")]
    assert ids
    r = client.post(f"/api/deid/jobs/{job_id}/confirm", json={"entity_ids": ids})
    assert r.status_code == 200


def test_confirm_requires_program_scan_ack(client, db, seeded_db):
    from docx import Document

    doc = Document()
    doc.add_paragraph("国家电投集团审计底稿内容。" * 8)
    buf = __import__("io").BytesIO()
    doc.save(buf)
    buf.seek(0)
    r = client.post(
        "/api/deid/jobs",
        files={"file": ("t.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    job_id = r.json()["id"]
    client.post(f"/api/deid/jobs/{job_id}/scan")
    client.post(f"/api/deid/jobs/{job_id}/semantic/skip")
    r = client.post(f"/api/deid/jobs/{job_id}/confirm", json={"entity_ids": []})
    assert "请先完成程序扫描" in r.text
